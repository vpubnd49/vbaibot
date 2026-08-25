#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Tuple

BASE = Path(__file__).resolve().parents[1]
REF = BASE / "references"
PROFILE = json.loads((REF / "review_profile.json").read_text(encoding="utf-8"))

def normalize_space(s: str) -> str:
    return re.sub(r"[ \t]+", " ", s or "")

def read_file(path: Path) -> Tuple[str, List[str]]:
    warnings: List[str] = []
    suf = path.suffix.lower()
    if suf in {".txt", ".md", ".csv", ".json", ".xml", ".html"}:
        return path.read_text(encoding="utf-8", errors="replace"), warnings

    if suf == ".docx":
        try:
            from docx import Document  # type: ignore
            doc = Document(str(path))
            chunks = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    chunks.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(chunks), warnings
        except Exception as e:
            warnings.append(f"Không đọc được DOCX bằng python-docx: {e}")
            return "", warnings

    if suf == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore
            reader = PdfReader(str(path))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
            if not text.strip():
                warnings.append("PDF không có text trích xuất được; có thể là bản scan. Cần OCR/vision.")
            return text, warnings
        except Exception as e:
            warnings.append(f"Không đọc được PDF bằng pypdf: {e}")
            return "", warnings

    warnings.append(f"Định dạng {suf or '(không phần mở rộng)'} chưa được script hỗ trợ.")
    return "", warnings

def issue(severity: str, category: str, message: str, evidence: str = "", suggestion: str = "") -> Dict[str, Any]:
    return {
        "severity": severity,
        "category": category,
        "message": message,
        "evidence": evidence,
        "suggestion": suggestion,
    }

def scan_text(text: str, label: str) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    low = text.lower()

    # Common typos.
    for wrong, right in PROFILE.get("common_typos", {}).items():
        if wrong.lower() in low:
            out.append(issue(
                "must_fix", "spelling",
                f"Phát hiện lỗi chính tả: “{wrong}”.",
                wrong,
                f"Sửa thành “{right}”."
            ))

    # Reporting logic: aggregation/reporting but explicit no separate reports.
    if all(term in low for term in ["tổng hợp", "báo cáo"]) and "không yêu cầu" in low and "báo cáo riêng" in low:
        out.append(issue(
            "should_fix", "reporting_logic",
            "Có dấu hiệu mâu thuẫn: giao tổng hợp/báo cáo nhưng đồng thời không yêu cầu báo cáo riêng.",
            "tổng hợp + báo cáo + không yêu cầu ... báo cáo riêng",
            "Làm rõ nguồn dữ liệu: chế độ báo cáo hiện hành và cơ chế yêu cầu cung cấp/cập nhật bổ sung khi thiếu."
        ))

    if "định kỳ" in low and "báo cáo" in low:
        out.append(issue(
            "verify", "reporting_logic",
            "Có sử dụng “định kỳ” đối với báo cáo. Cần xác minh kỳ báo cáo/căn cứ giao báo cáo định kỳ.",
            "định kỳ ... báo cáo",
            "Nếu không có kỳ báo cáo cụ thể, cân nhắc bỏ “định kỳ” hoặc dẫn căn cứ."
        ))

    # Direct vs advise reporting can change the workflow.
    if "báo cáo trực tiếp" in low and "tham mưu" in low and "báo cáo" in low:
        out.append(issue(
            "verify", "task_assignment",
            "Cùng tài liệu có cả “báo cáo trực tiếp” và “tham mưu ... báo cáo”. Cần kiểm tra có hai luồng báo cáo khác nhau hay biên tập làm đổi bản chất.",
            "báo cáo trực tiếp / tham mưu ... báo cáo",
            "Đối chiếu văn bản nguồn để xác định đúng chủ thể ký và nơi nhận."
        ))

    # Similar CTQH numbers: don't auto-correct.
    ctqh = sorted(set(re.findall(r"\b(\d{1,3})/CTQH\b", text, flags=re.I)))
    if len(ctqh) >= 2:
        out.append(issue(
            "verify", "citation",
            f"Phát hiện nhiều số hiệu CTQH gần nhau: {', '.join(x + '/CTQH' for x in ctqh)}.",
            ", ".join(x + "/CTQH" for x in ctqh),
            "Không tự sửa số hiệu. Đọc từng văn bản gốc để xác định chúng có phải các văn bản độc lập."
        ))

    # Potential month/year chronology inconsistency in header.
    header_match = re.search(
        r"(?:Lâm Đồng|Hà Nội|[\wÀ-ỹ ]+),\s*ngày\s*(?:\d{1,2})?\s*tháng\s*(\d{1,2})\s*năm\s*(20\d{2})",
        text, flags=re.I
    )
    if header_match:
        h_month, h_year = int(header_match.group(1)), int(header_match.group(2))
        dates = [(int(d), int(m), int(y), raw) for raw,d,m,y in re.findall(
            r"((\d{1,2})/(\d{1,2})/(20\d{2}))", text
        )]
        later = [raw for d,m,y,raw in dates if (y, m) > (h_year, h_month)]
        if later:
            out.append(issue(
                "verify", "chronology",
                "Có ngày được viện dẫn muộn hơn tháng/năm thể hiện ở dòng địa danh - ngày tháng của văn bản.",
                ", ".join(later[:5]),
                "Kiểm tra ngày ký dự kiến; nếu văn bản chưa ký thì có thể chỉ là placeholder, nếu đã ký thì phải sửa."
            ))

    # Repeated punctuation / spacing.
    for pat, msg, ev in [
        (r"\s+[,.;:]", "Có khoảng trắng thừa trước dấu câu.", " … ,/;/./:"),
        (r",{2,}|;{2,}|\.{3,}", "Có dấu câu lặp bất thường.", "dấu câu lặp"),
        (r"\bỦy ban nhân dân\s*\(UBND\)\s*tỉnh.*\bUBND tỉnh\b", "Kiểm tra tính nhất quán của viết tắt UBND trong cùng đoạn.", "Ủy ban nhân dân (UBND) / UBND"),
    ]:
        if re.search(pat, text, flags=re.I):
            out.append(issue("editorial", "punctuation" if "dấu" in msg or "khoảng" in msg else "terminology", msg, ev, "Biên tập lại nếu cần."))

    return out

def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()

def main() -> None:
    ap = argparse.ArgumentParser(description="Preflight rà soát văn bản hành chính Việt Nam.")
    ap.add_argument("--file", action="append", default=[], help="Đường dẫn file; có thể lặp nhiều lần.")
    ap.add_argument("--text", help="Nội dung text trực tiếp.")
    args = ap.parse_args()

    records: List[Dict[str, Any]] = []
    all_texts: List[Tuple[str, str]] = []
    global_warnings: List[str] = []

    for f in args.file:
        p = Path(f)
        if not p.exists():
            global_warnings.append(f"Không tìm thấy file: {p}")
            continue
        text, warnings = read_file(p)
        global_warnings.extend(f"{p.name}: {w}" for w in warnings)
        issues = scan_text(text, p.name) if text else []
        records.append({
            "file": str(p),
            "sha256": sha256(p),
            "chars_extracted": len(text),
            "issues": issues,
        })
        if text:
            all_texts.append((p.name, text))

    if args.text:
        issues = scan_text(args.text, "inline_text")
        records.append({
            "file": None,
            "label": "inline_text",
            "chars_extracted": len(args.text),
            "issues": issues,
        })
        all_texts.append(("inline_text", args.text))

    # Cross-file signal: different CTQH references across the corpus.
    corpus = "\n".join(t for _, t in all_texts)
    corpus_ctqh = sorted(set(re.findall(r"\b(\d{1,3})/CTQH\b", corpus, flags=re.I)))
    cross_issues: List[Dict[str, Any]] = []
    if len(all_texts) > 1 and len(corpus_ctqh) >= 2:
        cross_issues.append(issue(
            "verify", "cross_document",
            f"Nhiều file trong hồ sơ nhắc các số hiệu CTQH khác nhau: {', '.join(x + '/CTQH' for x in corpus_ctqh)}.",
            ", ".join(corpus_ctqh),
            "Đối chiếu từng văn bản gốc; không tự đồng nhất các số hiệu."
        ))

    severity_order = {"must_fix": 0, "verify": 1, "should_fix": 2, "editorial": 3}
    counts = {k: 0 for k in severity_order}
    for r in records:
        for i in r["issues"]:
            counts[i["severity"]] = counts.get(i["severity"], 0) + 1
    for i in cross_issues:
        counts[i["severity"]] = counts.get(i["severity"], 0) + 1

    if counts["must_fix"]:
        status = "needs_revision"
    elif counts["verify"]:
        status = "needs_verification"
    elif counts["should_fix"] or counts["editorial"]:
        status = "needs_revision"
    else:
        status = "clean"

    result = {
        "status": status,
        "note": "Kết quả script là preflight cơ học, không thay thế rà soát ngữ cảnh và kiểm tra trực quan.",
        "counts": counts,
        "records": records,
        "cross_document_issues": cross_issues,
        "warnings": global_warnings,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    main()
